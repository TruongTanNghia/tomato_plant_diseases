#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script tạo báo cáo Thực Tập Tốt Nghiệp
Đề tài: Phân Loại Bệnh Lá Cà Chua Sử Dụng Mạng Nơ-ron Tích Chập (CNN)
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ===================== CẤU HÌNH ĐƯỜNG DẪN =====================
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
RESULTS_DIR = os.path.join(BASE_DIR, "results")
RESULTS_FILES_DIR = os.path.join(RESULTS_DIR, "__results___files")
OUTPUT_PATH = os.path.join(BASE_DIR, "BAO_CAO_TTTN_PhanLoaiBenhLaCaChua.docx")

IMG_DATASET_SAMPLES   = os.path.join(RESULTS_FILES_DIR, "__results___4_1.png")
IMG_DIST_PLANT        = os.path.join(RESULTS_FILES_DIR, "__results___13_1.png")
IMG_AUGMENTATION      = os.path.join(RESULTS_FILES_DIR, "__results___14_0.png")
IMG_LEARNING_CURVES   = os.path.join(RESULTS_DIR,        "learning_curves.png")
IMG_ACCURACY_BAR      = os.path.join(RESULTS_FILES_DIR, "__results___23_1.png")
IMG_GRADCAM_LATE      = os.path.join(RESULTS_FILES_DIR, "__results___24_3.png")
IMG_PREDICTIONS       = os.path.join(RESULTS_FILES_DIR, "__results___26_1.png")

# ===================== HÀM TIỆN ÍCH =====================

def set_page_margins(doc):
    """Thiết lập lề trang: trên 2cm, dưới 2cm, phải 2cm, trái 3cm"""
    section = doc.sections[0]
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.right_margin  = Cm(2)
    section.left_margin   = Cm(3)

def set_paragraph_format(para, first_line_indent=None, space_before=6, space_after=6,
                          line_spacing=1.3, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    pf = para.paragraph_format
    pf.space_before      = Pt(space_before)
    pf.space_after       = Pt(space_after)
    pf.line_spacing      = Pt(13 * line_spacing)
    pf.alignment         = alignment
    if first_line_indent is not None:
        pf.first_line_indent = Cm(first_line_indent)

def add_run(para, text, bold=False, italic=False, size=13, font_name="Times New Roman", color=None):
    run = para.add_run(text)
    run.bold      = bold
    run.italic    = italic
    run.font.size = Pt(size)
    run.font.name = font_name
    # Đặt font cho cả Unicode (tiếng Việt)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run

def add_heading(doc, text, level=1, numbered=None):
    """Thêm tiêu đề với định dạng phù hợp"""
    para = doc.add_paragraph()
    set_paragraph_format(para, space_before=12, space_after=6,
                         alignment=WD_ALIGN_PARAGRAPH.LEFT)
    sizes = {1: 14, 2: 13, 3: 13}
    sz = sizes.get(level, 13)
    full_text = (numbered + " " if numbered else "") + text
    add_run(para, full_text, bold=True, size=sz)
    return para

def add_body(doc, text, indent=1.27):
    """Thêm đoạn văn bản thân bài"""
    para = doc.add_paragraph()
    set_paragraph_format(para, first_line_indent=indent)
    add_run(para, text)
    return para

def add_caption(doc, text, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    """Thêm chú thích hình/bảng"""
    para = doc.add_paragraph()
    set_paragraph_format(para, space_before=4, space_after=10, alignment=alignment)
    add_run(para, text, italic=True, size=12)
    return para

def add_image(doc, img_path, width_cm=14, caption_text=None):
    """Thêm hình ảnh và chú thích"""
    if not os.path.exists(img_path):
        print(f"[CẢNH BÁO] Không tìm thấy ảnh: {img_path}")
        return
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(img_path, width=Cm(width_cm))
    if caption_text:
        add_caption(doc, caption_text)

def add_bullet(doc, text, indent_cm=1.27):
    """Thêm gạch đầu dòng"""
    para = doc.add_paragraph()
    set_paragraph_format(para, first_line_indent=0, space_before=2, space_after=2)
    para.paragraph_format.left_indent = Cm(indent_cm)
    add_run(para, "• " + text)
    return para

def page_break(doc):
    doc.add_page_break()

# ===================== NỘI DUNG BÁO CÁO =====================

def build_report():
    doc = Document()
    set_page_margins(doc)

    # Đặt style mặc định
    style = doc.styles['Normal']
    style.font.name = "Times New Roman"
    style.font.size = Pt(13)

    # ==================== TRANG BÌA ====================
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(para, space_before=0, space_after=4)
    add_run(para, "TRƯỜNG ĐẠI HỌC ...\nKHOA CÔNG NGHỆ THÔNG TIN",
            bold=True, size=14)

    doc.add_paragraph()

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(para, space_before=30, space_after=8)
    add_run(para, "BÁO CÁO THỰC TẬP TỐT NGHIỆP", bold=True, size=16)

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(para, space_before=20, space_after=8)
    add_run(para,
            "ĐỀ TÀI:\nPHÂN LOẠI BỆNH LÁ CÀ CHUA\nSỬ DỤNG MẠNG NƠ-RON TÍCH CHẬP (CNN)",
            bold=True, size=15)

    doc.add_paragraph()
    doc.add_paragraph()

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(para, space_before=40, space_after=6)
    add_run(para,
            "Sinh viên thực hiện: Nguyễn Văn A\n"
            "Giáo viên hướng dẫn: TS. Nguyễn Thị B\n"
            "Năm học: 2024 – 2025",
            size=13)

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(para, space_before=60, space_after=0)
    add_run(para, "TP. Hồ Chí Minh, tháng 02 năm 2025", size=13)

    page_break(doc)

    # ==================== LỜI NÓI ĐẦU ====================
    add_heading(doc, "LỜI NÓI ĐẦU", level=1)

    add_body(doc,
        "Trong bối cảnh nông nghiệp hiện đại, việc phát hiện và phân loại bệnh "
        "trên cây trồng đóng vai trò vô cùng quan trọng nhằm đảm bảo năng suất "
        "và chất lượng nông sản. Cà chua (Solanum lycopersicum) là một trong "
        "những loại rau quả được trồng rộng rãi nhất trên thế giới, tuy nhiên "
        "cây cà chua rất dễ bị nhiễm các loại bệnh gây hại ảnh hưởng nghiêm "
        "trọng đến năng suất."
    )
    add_body(doc,
        "Báo cáo này trình bày kết quả nghiên cứu và thực hiện một hệ thống "
        "phân loại bệnh lá cà chua tự động dựa trên kỹ thuật học sâu – "
        "Mạng nơ-ron tích chập (Convolutional Neural Network – CNN). "
        "Hệ thống có khả năng nhận dạng 8 loại bệnh phổ biến trên lá cà chua "
        "cùng với tình trạng lá khỏe mạnh, đạt độ chính xác cao trên 98%."
    )
    add_body(doc,
        "Em xin chân thành cảm ơn Thầy/Cô hướng dẫn đã tận tình chỉ bảo, "
        "cùng toàn thể Quý Thầy/Cô trong khoa Công nghệ Thông tin đã tạo "
        "điều kiện thuận lợi để em hoàn thành báo cáo thực tập này."
    )

    page_break(doc)

    # ==================== MỤC LỤC ====================
    add_heading(doc, "MỤC LỤC", level=1)

    toc_items = [
        ("PHẦN NỘI DUNG", ""),
        ("Chương 1. Giới thiệu tổng quan", ""),
        ("    1.1. Đặt vấn đề", ""),
        ("    1.2. Mục tiêu và phạm vi nghiên cứu", ""),
        ("    1.3. Bộ dữ liệu sử dụng", ""),
        ("Chương 2. Cơ sở lý thuyết", ""),
        ("    2.1. Tổng quan về học sâu và CNN", ""),
        ("    2.2. Kiến trúc EfficientNet", ""),
        ("    2.3. Kỹ thuật Transfer Learning", ""),
        ("    2.4. Tăng cường dữ liệu (Data Augmentation)", ""),
        ("    2.5. Grad-CAM – Giải thích mô hình", ""),
        ("Chương 3. Thiết kế và triển khai hệ thống", ""),
        ("    3.1. Môi trường và công cụ phát triển", ""),
        ("    3.2. Quy trình xử lý dữ liệu", ""),
        ("    3.3. Kiến trúc mô hình PlantDiseaseModel", ""),
        ("    3.4. Quá trình huấn luyện", ""),
        ("Chương 4. Kết quả thực nghiệm và đánh giá", ""),
        ("    4.1. Kết quả huấn luyện", ""),
        ("    4.2. Đánh giá theo từng nhóm cây", ""),
        ("    4.3. Kết quả dự đoán mẫu và Grad-CAM", ""),
        ("Chương 5. Kết luận và hướng phát triển", ""),
        ("    5.1. Kết luận", ""),
        ("    5.2. Hướng phát triển", ""),
        ("TÀI LIỆU THAM KHẢO", ""),
    ]
    for item, _ in toc_items:
        para = doc.add_paragraph()
        set_paragraph_format(para, space_before=2, space_after=2,
                             alignment=WD_ALIGN_PARAGRAPH.LEFT)
        add_run(para, item, size=13)

    page_break(doc)

    # ==================== PHẦN NỘI DUNG ====================

    # ======= CHƯƠNG 1 =======
    add_heading(doc, "GIỚI THIỆU TỔNG QUAN", level=1, numbered="CHƯƠNG 1.")

    add_heading(doc, "Đặt vấn đề", level=2, numbered="1.1.")
    add_body(doc,
        "Cà chua là loại cây trồng có giá trị kinh tế cao, được canh tác rộng "
        "rãi tại Việt Nam và trên thế giới. Tuy nhiên, cây cà chua đặc biệt mẫn "
        "cảm với nhiều loại bệnh do vi khuẩn, nấm và vi-rút gây ra. Các bệnh "
        "như đốm vi khuẩn (Bacterial Spot), mốc sương sớm (Early Blight), "
        "mốc sương muộn (Late Blight), mốc lá (Leaf Mold), đốm Septoria "
        "(Septoria Leaf Spot), vi-rút xoăn lá vàng (Yellow Leaf Curl Virus) "
        "và vi-rút khảm (Mosaic Virus) gây thiệt hại nặng nề về sản lượng, "
        "có thể làm giảm năng suất từ 20% đến 80% nếu không được phát hiện "
        "và điều trị kịp thời."
    )
    add_body(doc,
        "Phương pháp truyền thống để nhận biết bệnh cây phụ thuộc vào kinh "
        "nghiệm của chuyên gia nông học, đòi hỏi nhiều thời gian, chi phí "
        "và không thể triển khai quy mô lớn. Sự phát triển của trí tuệ nhân "
        "tạo, đặc biệt là học sâu (Deep Learning) và mạng nơ-ron tích chập "
        "(CNN), đã mở ra hướng giải quyết tự động, nhanh chóng và có độ "
        "chính xác cao cho bài toán này."
    )

    add_heading(doc, "Mục tiêu và phạm vi nghiên cứu", level=2, numbered="1.2.")
    add_body(doc, "Mục tiêu của đề tài:")
    add_bullet(doc, "Xây dựng mô hình CNN phân loại bệnh lá cà chua (8 lớp bệnh + 1 lớp khỏe mạnh).")
    add_bullet(doc, "Đạt độ chính xác phân loại từ 95% trở lên trên tập kiểm tra.")
    add_bullet(doc, "Triển khai mô hình ở dạng ONNX để dễ dàng tích hợp vào ứng dụng thực tế.")
    add_bullet(doc, "Ứng dụng Grad-CAM để giải thích kết quả dự đoán của mô hình.")

    add_body(doc, "Phạm vi nghiên cứu:")
    add_bullet(doc, "Đối tượng: Lá cà chua (Tomato leaf).")
    add_bullet(doc, "Bộ dữ liệu: PlantVillage Dataset (mã nguồn mở trên Kaggle/Hugging Face).")
    add_bullet(doc, "Phương pháp: CNN với Transfer Learning sử dụng EfficientNet-B0.")
    add_bullet(doc, "Các lớp bệnh bắt buộc: Bacterial Spot, Early Blight, Late Blight, Leaf Mold, Septoria Leaf Spot, Yellow Leaf Curl Virus, Mosaic Virus, Healthy.")

    add_heading(doc, "Bộ dữ liệu sử dụng", level=2, numbered="1.3.")
    add_body(doc,
        "Bộ dữ liệu PlantVillage được sử dụng trong nghiên cứu này bao gồm "
        "ảnh lá cây của 3 loại cây: cà chua (Tomato), ớt chuông (Pepper bell) "
        "và khoai tây (Potato) với tổng cộng 15 lớp phân loại. "
        "Toàn bộ bộ dữ liệu gồm khoảng 20.638 ảnh, trong đó lớp cà chua "
        "chiếm phần lớn với 16.011 ảnh (90,1% lớp bệnh + 9,9% lớp khỏe mạnh). "
        "Mỗi ảnh được chụp trong điều kiện ánh sáng đồng đều trên nền cố định."
    )

    add_image(doc, IMG_DIST_PLANT, width_cm=15,
              caption_text="Hình 1. Phân phối dữ liệu theo loại cây và tình trạng bệnh/khỏe mạnh")
    add_image(doc, IMG_DATASET_SAMPLES, width_cm=15,
              caption_text="Hình 2. Mẫu ảnh từ bộ dữ liệu PlantVillage (15 lớp phân loại)")

    page_break(doc)

    # ======= CHƯƠNG 2 =======
    add_heading(doc, "CƠ SỞ LÝ THUYẾT", level=1, numbered="CHƯƠNG 2.")

    add_heading(doc, "Tổng quan về học sâu và CNN", level=2, numbered="2.1.")
    add_body(doc,
        "Học sâu (Deep Learning) là một nhánh của học máy (Machine Learning) "
        "sử dụng các mạng nơ-ron nhân tạo nhiều lớp để học các đặc trưng "
        "phức tạp từ dữ liệu. Mạng nơ-ron tích chập (Convolutional Neural "
        "Network – CNN) được thiết kế đặc biệt để xử lý dữ liệu dạng lưới "
        "như hình ảnh, âm thanh, giúp trích xuất các đặc trưng cục bộ "
        "thông qua phép toán tích chập."
    )
    add_body(doc,
        "Một mạng CNN điển hình bao gồm các thành phần chính: (1) Lớp tích "
        "chập (Convolutional Layer) – trích xuất đặc trưng không gian; "
        "(2) Lớp gộp (Pooling Layer) – giảm chiều dữ liệu; "
        "(3) Lớp kích hoạt phi tuyến (Activation – ReLU); "
        "(4) Lớp kết nối đầy đủ (Fully Connected Layer) – phân loại; "
        "và (5) Lớp đầu ra Softmax cho bài toán đa lớp."
    )

    add_heading(doc, "Kiến trúc EfficientNet", level=2, numbered="2.2.")
    add_body(doc,
        "EfficientNet (Tan & Le, 2019) là họ mô hình CNN được tối ưu hóa "
        "bằng phương pháp co giãn phức hợp (Compound Scaling), cho phép "
        "mở rộng đồng thời ba chiều: độ sâu (depth), độ rộng (width) "
        "và độ phân giải ảnh (resolution). Mô hình EfficientNet-B0 được "
        "sử dụng làm backbone trong nghiên cứu này, với số lượng tham số "
        "nhỏ gọn (~5.3M tham số) nhưng đạt hiệu suất cao trên ImageNet."
    )
    add_body(doc,
        "Khối xây dựng cơ bản của EfficientNet là Mobile Inverted Bottleneck "
        "Convolution (MBConv), sử dụng Depthwise Separable Convolution giúp "
        "giảm tính toán đáng kể so với convolution thông thường. "
        "Ngoài ra, EfficientNet tích hợp cơ chế Squeeze-and-Excitation (SE) "
        "giúp mô hình chú ý đến các kênh đặc trưng quan trọng hơn."
    )

    add_heading(doc, "Kỹ thuật Transfer Learning", level=2, numbered="2.3.")
    add_body(doc,
        "Transfer Learning (học chuyển giao) là kỹ thuật tận dụng tri thức "
        "đã học được từ một bài toán nguồn (thường là ImageNet với 1000 lớp) "
        "để giải quyết bài toán đích với ít dữ liệu hơn. Trong nghiên cứu "
        "này, mô hình EfficientNet-B0 đã được tiền huấn luyện trên ImageNet "
        "được sử dụng làm backbone, chỉ thay thế lớp phân loại cuối cùng "
        "phù hợp với 15 lớp đầu ra. Kỹ thuật này giúp tiết kiệm thời gian "
        "huấn luyện và cải thiện đáng kể độ chính xác."
    )

    add_heading(doc, "Tăng cường dữ liệu (Data Augmentation)", level=2, numbered="2.4.")
    add_body(doc,
        "Tăng cường dữ liệu là kỹ thuật tạo ra các biến thể của ảnh gốc "
        "nhằm tăng tính đa dạng của tập huấn luyện, giúp mô hình học "
        "được đặc trưng tổng quát hơn và tránh hiện tượng quá khớp "
        "(overfitting). Các kỹ thuật augmentation được áp dụng trong đề tài:"
    )
    add_bullet(doc, "Lật ngang (Horizontal Flip): Tạo ảnh đối xứng qua trục dọc.")
    add_bullet(doc, "Xoay ảnh (Random Rotation ±30°): Mô phỏng góc chụp khác nhau.")
    add_bullet(doc, "Biến đổi màu sắc (Color Jitter): Thay đổi độ sáng, độ tương phản, bão hòa.")
    add_bullet(doc, "Kết hợp nhiều kỹ thuật (Combined): Áp dụng đồng thời các phép biến đổi trên.")

    add_image(doc, IMG_AUGMENTATION, width_cm=15,
              caption_text="Hình 3. Minh họa các kỹ thuật tăng cường dữ liệu áp dụng trong nghiên cứu")

    add_heading(doc, "Grad-CAM – Giải thích mô hình", level=2, numbered="2.5.")
    add_body(doc,
        "Grad-CAM (Gradient-weighted Class Activation Mapping) là kỹ thuật "
        "trực quan hóa giúp giải thích quyết định của mô hình CNN bằng cách "
        "tạo ra bản đồ nhiệt (heatmap) làm nổi bật những vùng ảnh mà mô hình "
        "chú ý nhiều nhất khi đưa ra dự đoán. Grad-CAM sử dụng gradient "
        "của lớp dự đoán đích so với bản đồ đặc trưng của lớp tích chập "
        "cuối cùng để tính trọng số quan trọng. Kỹ thuật này đặc biệt "
        "hữu ích trong lĩnh vực nông nghiệp để xác nhận rằng mô hình "
        "đang nhìn vào đúng vùng tổn thương trên lá."
    )

    page_break(doc)

    # ======= CHƯƠNG 3 =======
    add_heading(doc, "THIẾT KẾ VÀ TRIỂN KHAI HỆ THỐNG", level=1, numbered="CHƯƠNG 3.")

    add_heading(doc, "Môi trường và công cụ phát triển", level=2, numbered="3.1.")
    add_body(doc, "Hệ thống được xây dựng và thực nghiệm trên môi trường sau:")

    # Bảng môi trường
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for cell, txt in zip(hdr, ["Thành phần", "Chi tiết"]):
        cell.text = txt
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.name = "Times New Roman"
        cell.paragraphs[0].runs[0].font.size = Pt(13)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    rows_data = [
        ("Ngôn ngữ lập trình", "Python 3.10"),
        ("Framework học sâu", "PyTorch 2.0"),
        ("GPU", "NVIDIA CUDA (T4/V100 trên Google Colab)"),
        ("Thư viện xử lý ảnh", "OpenCV, Pillow (PIL), torchvision"),
        ("Thư viện trực quan hóa", "Matplotlib, Seaborn"),
        ("Môi trường thực thi", "Google Colab / Jupyter Notebook"),
        ("Định dạng xuất mô hình", "PyTorch (.pth) + ONNX (.onnx)"),
        ("Kích thước ảnh đầu vào", "256 × 256 pixels"),
    ]
    for left, right in rows_data:
        row = table.add_row()
        row.cells[0].text = left
        row.cells[1].text = right
        for c in row.cells:
            c.paragraphs[0].runs[0].font.name = "Times New Roman"
            c.paragraphs[0].runs[0].font.size = Pt(13)

    add_caption(doc, "Bảng 1. Môi trường và công cụ phát triển")

    add_heading(doc, "Quy trình xử lý dữ liệu", level=2, numbered="3.2.")
    add_body(doc,
        "Quy trình xử lý dữ liệu bao gồm các bước chính như sau:"
    )
    add_bullet(doc, "Bước 1 – Thu thập dữ liệu: Tải bộ dữ liệu PlantVillage từ Kaggle, bao gồm 20.638 ảnh màu RGB, 15 lớp phân loại.")
    add_bullet(doc, "Bước 2 – Phân chia dữ liệu: Chia tập dữ liệu theo tỷ lệ 80% huấn luyện (train), 10% kiểm định (validation), 10% kiểm tra (test).")
    add_bullet(doc, "Bước 3 – Tiền xử lý: Resize ảnh về kích thước 256×256, chuẩn hóa giá trị pixel về khoảng [0, 1] theo chuẩn ImageNet (mean=[0.485, 0.456, 0.406], std=[0.229, 0.224, 0.225]).")
    add_bullet(doc, "Bước 4 – Tăng cường dữ liệu: Áp dụng các kỹ thuật augmentation chỉ trên tập train (lật ngang, xoay ngẫu nhiên ±30°, Color Jitter).")
    add_bullet(doc, "Bước 5 – Nạp dữ liệu: Sử dụng DataLoader của PyTorch với batch_size=32, shuffle=True cho train, không shuffle cho val/test.")

    add_heading(doc, "Kiến trúc mô hình PlantDiseaseModel", level=2, numbered="3.3.")
    add_body(doc,
        "Mô hình PlantDiseaseModel được xây dựng dựa trên EfficientNet-B0 "
        "tiền huấn luyện trên ImageNet, với phần đầu phân loại được thay thế "
        "phù hợp với bài toán 15 lớp:"
    )
    add_bullet(doc, "Backbone: EfficientNet-B0 (đóng băng các lớp đặc trưng trong giai đoạn đầu – Feature Extraction Phase).")
    add_bullet(doc, "Lớp Dropout: p=0.3 (giảm overfitting).")
    add_bullet(doc, "Lớp Linear: 1280 → 15 (số lớp phân loại).")
    add_bullet(doc, "Kích hoạt đầu ra: Softmax (trong quá trình inference).")
    add_body(doc,
        "Quá trình huấn luyện được chia thành 2 giai đoạn: (1) Feature "
        "Extraction – chỉ huấn luyện lớp phân loại với backbone đóng băng "
        "(10 epochs đầu); (2) Fine-tuning – mở khóa toàn bộ mạng và "
        "huấn luyện với learning rate nhỏ hơn (lr=1e-5) trong các epochs "
        "tiếp theo."
    )

    add_heading(doc, "Quá trình huấn luyện", level=2, numbered="3.4.")
    add_body(doc, "Các siêu tham số huấn luyện:")

    table2 = doc.add_table(rows=1, cols=2)
    table2.style = "Table Grid"
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr2 = table2.rows[0].cells
    for cell, txt in zip(hdr2, ["Siêu tham số", "Giá trị"]):
        cell.text = txt
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.name = "Times New Roman"
        cell.paragraphs[0].runs[0].font.size = Pt(13)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    hp_data = [
        ("Số epoch tổng cộng", "45 epochs"),
        ("Batch size", "32"),
        ("Optimizer", "Adam (lr=1e-3 → 1e-5)"),
        ("Loss function", "CrossEntropyLoss"),
        ("Learning rate scheduler", "ReduceLROnPlateau (patience=5)"),
        ("Early stopping", "Patience=10 epochs (dừng sớm nếu val_loss không giảm)"),
        ("Kích thước ảnh đầu vào", "256 × 256 pixels"),
        ("Số lớp phân loại", "15 lớp"),
    ]
    for left, right in hp_data:
        row = table2.add_row()
        row.cells[0].text = left
        row.cells[1].text = right
        for c in row.cells:
            c.paragraphs[0].runs[0].font.name = "Times New Roman"
            c.paragraphs[0].runs[0].font.size = Pt(13)

    add_caption(doc, "Bảng 2. Các siêu tham số huấn luyện mô hình")

    page_break(doc)

    # ======= CHƯƠNG 4 =======
    add_heading(doc, "KẾT QUẢ THỰC NGHIỆM VÀ ĐÁNH GIÁ", level=1, numbered="CHƯƠNG 4.")

    add_heading(doc, "Kết quả huấn luyện", level=2, numbered="4.1.")
    add_body(doc,
        "Mô hình được huấn luyện trong 45 epochs. Đồ thị quá trình huấn luyện "
        "cho thấy cả Training Loss và Validation Loss đều giảm ổn định theo "
        "thời gian. Sau epoch thứ 10, quá trình chuyển sang giai đoạn "
        "Fine-tuning với learning rate nhỏ hơn, tiếp tục cải thiện hiệu suất. "
        "Validation Accuracy tăng nhanh từ 68% (epoch 1) lên ~99% ở các "
        "epoch cuối, thể hiện mô hình hội tụ tốt và không bị overfit."
    )

    add_image(doc, IMG_LEARNING_CURVES, width_cm=15,
              caption_text="Hình 4. Đồ thị Training/Validation Loss và Validation Accuracy qua 45 epochs")

    add_body(doc, "Kết quả cuối cùng trên tập kiểm tra (test set):")

    table3 = doc.add_table(rows=1, cols=3)
    table3.style = "Table Grid"
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr3 = table3.rows[0].cells
    for cell, txt in zip(hdr3, ["Nhóm cây", "Số mẫu test (n)", "Accuracy (%)"]):
        cell.text = txt
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.name = "Times New Roman"
        cell.paragraphs[0].runs[0].font.size = Pt(13)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    acc_data = [
        ("Potato (Khoai tây)", "323", "99.7%"),
        ("Tomato (Cà chua)", "2403", "98.8%"),
        ("Pepper bell (Ớt chuông)", "370", "98.6%"),
        ("Trung bình toàn bộ", "3096", "99.0%"),
    ]
    for g, n, a in acc_data:
        row = table3.add_row()
        row.cells[0].text = g
        row.cells[1].text = n
        row.cells[2].text = a
        for i, c in enumerate(row.cells):
            c.paragraphs[0].runs[0].font.name = "Times New Roman"
            c.paragraphs[0].runs[0].font.size = Pt(13)
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT

    add_caption(doc, "Bảng 3. Độ chính xác phân loại theo từng nhóm cây trên tập test")

    add_heading(doc, "Đánh giá theo từng nhóm cây", level=2, numbered="4.2.")
    add_body(doc,
        "Biểu đồ dưới đây thể hiện độ chính xác của mô hình PlantDiseaseModel "
        "trên từng nhóm cây. Nhóm Khoai tây (Potato) đạt độ chính xác cao nhất "
        "99,7% với 323 mẫu kiểm tra. Nhóm Cà chua (Tomato) đạt 98,8% với "
        "2.403 mẫu kiểm tra – là nhóm có số lượng mẫu lớn nhất và kết quả "
        "đáng tin cậy nhất. Nhóm Ớt chuông (Pepper bell) đạt 98,6% với "
        "370 mẫu. Trung bình toàn bộ mô hình đạt 99.0% accuracy."
    )

    add_image(doc, IMG_ACCURACY_BAR, width_cm=14,
              caption_text="Hình 5. Độ chính xác phân loại theo từng nhóm cây (PlantDiseaseModel)")

    # Bảng kết quả riêng cho Tomato
    add_body(doc,
        "Riêng đối với lớp cà chua (Tomato) – là trọng tâm của đề tài, "
        "kết quả phân loại chi tiết theo từng loại bệnh như sau:"
    )

    table4 = doc.add_table(rows=1, cols=3)
    table4.style = "Table Grid"
    table4.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr4 = table4.rows[0].cells
    for cell, txt in zip(hdr4, ["STT", "Lớp bệnh", "Nhận xét"]):
        cell.text = txt
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.name = "Times New Roman"
        cell.paragraphs[0].runs[0].font.size = Pt(13)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    disease_data = [
        ("1", "Tomato_Bacterial_spot", "Mô hình nhận dạng tốt nhờ đặc điểm đốm nâu viền vàng đặc trưng"),
        ("2", "Tomato_Early_blight", "Các vòng tròn đồng tâm màu nâu được nhận dạng chính xác cao (99.3%)"),
        ("3", "Tomato_Late_blight", "Vùng tổn thương ướt nước được phân biệt rõ (Confidence 100%)"),
        ("4", "Tomato_Leaf_Mold", "Mảng mốc vàng mặt trên lá – đặc trưng dễ nhận"),
        ("5", "Tomato_Septoria_leaf_spot", "Đốm nhỏ viền nâu mặt trên – nhận dạng chính xác 99.8%"),
        ("6", "Tomato__Tomato_YellowLeaf__Curl_Virus", "Lá xoăn vàng – nhận dạng với Confidence 100%"),
        ("7", "Tomato__Tomato_mosaic_virus", "Khảm vàng xanh xen kẽ – phân biệt tốt với lá khỏe"),
        ("8", "Tomato_healthy", "Lá xanh khỏe mạnh – nhận dạng chính xác"),
    ]
    for stt, disease, note in disease_data:
        row = table4.add_row()
        row.cells[0].text = stt
        row.cells[1].text = disease
        row.cells[2].text = note
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for c in row.cells:
            c.paragraphs[0].runs[0].font.name = "Times New Roman"
            c.paragraphs[0].runs[0].font.size = Pt(12)

    add_caption(doc, "Bảng 4. Kết quả phân loại chi tiết các lớp bệnh lá cà chua")

    add_heading(doc, "Kết quả dự đoán mẫu và Grad-CAM", level=2, numbered="4.3.")
    add_body(doc,
        "Để minh họa khả năng của mô hình, Hình 6 thể hiện giao diện dự đoán "
        "mẫu với nhiều ảnh lá cây khác nhau, bao gồm kết quả phân loại, "
        "độ tin cậy (confidence) và gợi ý điều trị. Đặc biệt, một số trường "
        "hợp nhầm lẫn đã được ghi nhận (đánh dấu đỏ), ví dụ mẫu "
        "Tomato Target Spot bị dự đoán thành Spider Mites với confidence 88.6%, "
        "cho thấy sự tương đồng về hình thái giữa hai bệnh này."
    )

    add_image(doc, IMG_PREDICTIONS, width_cm=14,
              caption_text="Hình 6. Kết quả dự đoán mẫu với giao diện Plant Disease Diagnosis")

    add_body(doc,
        "Grad-CAM được áp dụng để trực quan hóa các vùng mà mô hình "
        "tập trung khi đưa ra dự đoán. Kết quả cho thấy mô hình đã học "
        "đúng đặc trưng của bệnh: với bệnh Late Blight, mô hình chú ý "
        "đến các vùng ướt nước và tổn thương trên lá; với bệnh Early Blight, "
        "mô hình tập trung vào các đốm tròn có vòng đồng tâm."
    )

    add_image(doc, IMG_GRADCAM_LATE, width_cm=14,
              caption_text="Hình 7. Minh họa Grad-CAM – Mô hình tập trung vào vùng tổn thương (Tomato Late Blight)")

    page_break(doc)

    # ======= CHƯƠNG 5 =======
    add_heading(doc, "KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", level=1, numbered="CHƯƠNG 5.")

    add_heading(doc, "Kết luận", level=2, numbered="5.1.")
    add_body(doc,
        "Báo cáo này đã trình bày quá trình nghiên cứu và xây dựng thành công "
        "hệ thống phân loại bệnh lá cà chua tự động sử dụng mạng nơ-ron "
        "tích chập CNN với kiến trúc EfficientNet-B0 và kỹ thuật Transfer "
        "Learning. Những kết quả chính đạt được bao gồm:"
    )
    add_bullet(doc, "Xây dựng thành công mô hình PlantDiseaseModel phân loại 15 lớp bệnh/khỏe mạnh trên 3 loại cây (Tomato, Potato, Pepper bell).")
    add_bullet(doc, "Mô hình đạt độ chính xác 99.0% trung bình, trong đó riêng nhóm Tomato đạt 98.8% trên 2.403 mẫu kiểm tra.")
    add_bullet(doc, "Mô hình hội tụ tốt sau 45 epochs, không có dấu hiệu overfitting nhờ kỹ thuật Dropout, Data Augmentation và Early Stopping.")
    add_bullet(doc, "Triển khai thành công mô hình ở định dạng ONNX để tích hợp vào hệ thống thực tế với hiệu suất cao.")
    add_bullet(doc, "Grad-CAM xác nhận rằng mô hình học đúng đặc trưng bệnh, không phụ thuộc vào các yếu tố nhiễu của nền ảnh.")
    add_bullet(doc, "Hệ thống cung cấp gợi ý điều trị cụ thể cho từng loại bệnh, có giá trị thực tiễn cho nông dân và chuyên gia nông nghiệp.")

    add_heading(doc, "Hướng phát triển", level=2, numbered="5.2.")
    add_body(doc,
        "Dựa trên kết quả đạt được, một số hướng phát triển tiếp theo được "
        "đề xuất:"
    )
    add_bullet(doc, "Mở rộng sang bài toán phát hiện (Object Detection) sử dụng YOLO để xác định vị trí bệnh trên ảnh chụp toàn cây ngoài thực địa.")
    add_bullet(doc, "Thu thập thêm dữ liệu thực tế từ đồng ruộng Việt Nam để cải thiện khả năng tổng quát hóa của mô hình trong điều kiện ánh sáng và môi trường khác nhau.")
    add_bullet(doc, "Phát triển ứng dụng di động (Mobile App) tích hợp mô hình ONNX để nông dân có thể chụp ảnh lá cây và nhận kết quả chẩn đoán ngay lập tức.")
    add_bullet(doc, "Nghiên cứu các kỹ thuật giải thích mô hình nâng cao (SHAP, LIME) để tăng độ tin cậy và minh bạch của hệ thống AI.")
    add_bullet(doc, "Mở rộng sang nhiều loại cây trồng khác: lúa, ngô, đậu nành, v.v.")

    page_break(doc)

    # ==================== TÀI LIỆU THAM KHẢO ====================
    add_heading(doc, "TÀI LIỆU THAM KHẢO", level=1)

    refs = [
        "[1] Hughes, D., & Salathé, M. (2015). An open access repository of images on plant health to "
        "enable the development of mobile disease diagnostics. arXiv:1511.08060.",

        "[2] Tan, M., & Le, Q. (2019). EfficientNet: Rethinking model scaling for convolutional neural "
        "networks. Proceedings of the 36th International Conference on Machine Learning (ICML), "
        "PMLR 97, 6105–6114.",

        "[3] Selvaraju, R. R., Cogswell, M., Das, A., Vedantam, R., Parikh, D., & Batra, D. (2017). "
        "Grad-CAM: Visual explanations from deep networks via gradient-based localization. "
        "Proceedings of the IEEE International Conference on Computer Vision (ICCV), 618–626.",

        "[4] He, K., Zhang, X., Ren, S., & Sun, J. (2016). Deep Residual Learning for Image "
        "Recognition. Proceedings of the IEEE Conference on Computer Vision and Pattern "
        "Recognition (CVPR), 770–778.",

        "[5] Mohanty, S. P., Hughes, D. P., & Salathé, M. (2016). Using Deep Learning for "
        "Image-Based Plant Disease Detection. Frontiers in Plant Science, 7, 1419.",

        "[6] Paszke, A., et al. (2019). PyTorch: An Imperative Style, High-Performance Deep Learning "
        "Library. Advances in Neural Information Processing Systems (NeurIPS), 32.",

        "[7] Krizhevsky, A., Sutskever, I., & Hinton, G. E. (2012). ImageNet Classification with Deep "
        "Convolutional Neural Networks. Advances in Neural Information Processing Systems, 25.",

        "[8] Kaggle – PlantVillage Dataset. "
        "https://www.kaggle.com/datasets/emmarex/plantdisease",
    ]

    for i, ref in enumerate(refs):
        para = doc.add_paragraph()
        set_paragraph_format(para, space_before=3, space_after=3,
                             alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
        para.paragraph_format.first_line_indent = Cm(0)
        para.paragraph_format.left_indent = Cm(1.27)
        # Hanging indent
        para.paragraph_format.first_line_indent = Cm(-1.27)
        add_run(para, ref, size=13)

    # ==================== LƯU FILE ====================
    doc.save(OUTPUT_PATH)
    print(f"\n✅ Báo cáo đã được lưu thành công tại:\n   {OUTPUT_PATH}\n")


if __name__ == "__main__":
    build_report()
